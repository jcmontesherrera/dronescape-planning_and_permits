"""
Generate a Trip Logistics Checklist from campaigns.db + tern_plots.db.

Outputs
-------
  Markdown  docs/checklists/<trip-id>-checklist.md   (Obsidian-ready, all info)
  DOCX      docs/checklists/<trip-id>-checklist.docx (filled from UTAS template)
  Maps      docs/checklists/maps/<trip-id>-overview.png
            docs/checklists/maps/<trip-id>-<property>.png

Usage
-----
  # Markdown only (default)
  python scripts/generate_checklist.py --trip-id 05-DRW-TOP-2027-06

  # Markdown + DOCX (uses TEMPLATE_DOCX from paths.py unless --template given)
  python scripts/generate_checklist.py --trip-id 05-DRW-TOP-2027-06 --docx

  # Everything: markdown + DOCX + maps
  python scripts/generate_checklist.py --trip-id 05-DRW-TOP-2027-06 --docx --maps

  # Custom template path
  python scripts/generate_checklist.py --trip-id 05-DRW-TOP-2027-06 --docx --maps `
      --template C:/path/to/Template.docx
"""

from __future__ import annotations

import argparse
import math
import re
import shutil
import sqlite3
import sys
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from dronescape_planning.paths import (
    ARD_STATE,
    CAMPAIGNS,
    DOCS_CHECKLISTS,
    TEMPLATE_DOCX,
    TERN_PLOTS,
)

# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


@dataclass
class ItineraryDayMeta:
    visit_date: str
    day_number: int | None
    day_name: str
    travel: str
    property_visited: str
    notes: str
    accommodation: str
    accom_link: str
    booking_status: str


@dataclass
class PlotRow:
    plot: str
    latitude: float
    longitude: float
    property: str
    site_location_name: str
    visit_date: Optional[str]
    access_notes: Optional[str]
    collected: int
    bioregion: Optional[str]
    mvg: Optional[str]
    state: str  # derived from plot prefix


@dataclass
class PropertyGroup:
    name: str
    tenure: Optional[str]
    jurisdiction: Optional[str]
    access_status: Optional[str]
    email_address: Optional[str]
    notes: Optional[str]
    plots: list[PlotRow] = field(default_factory=list)

    @property
    def centroid_lat(self) -> float:
        return sum(p.latitude for p in self.plots) / len(self.plots)

    @property
    def centroid_lon(self) -> float:
        return sum(p.longitude for p in self.plots) / len(self.plots)

    @property
    def plot_ids(self) -> str:
        return ", ".join(p.plot for p in self.plots)

    @property
    def bioregions(self) -> str:
        seen: list[str] = []
        for p in self.plots:
            if p.bioregion and p.bioregion not in seen:
                seen.append(p.bioregion)
        return ", ".join(seen) or "(unknown)"

    @property
    def mvgs(self) -> str:
        seen: list[str] = []
        for p in self.plots:
            if p.mvg and p.mvg not in seen:
                seen.append(p.mvg)
        return ", ".join(seen) or "(unknown)"

    @property
    def state_codes(self) -> str:
        codes = sorted({p.state for p in self.plots})
        return ", ".join(codes)


@dataclass
class TripData:
    trip_id: str
    route_origin: str
    route_dest: str
    start_date: str
    end_date: str
    season: Optional[str]
    enso_phase: Optional[str]
    notes: Optional[str]
    properties: list[PropertyGroup] = field(default_factory=list)
    itinerary_days_meta: list[ItineraryDayMeta] = field(default_factory=list)

    @property
    def itinerary_imported(self) -> bool:
        return bool(self.itinerary_days_meta)

    @property
    def all_plots(self) -> list[PlotRow]:
        return [p for grp in self.properties for p in grp.plots]

    @property
    def location_names(self) -> str:
        names: list[str] = []
        for grp in self.properties:
            if grp.name not in names:
                names.append(grp.name)
        return ", ".join(names)

    @property
    def states(self) -> str:
        codes: list[str] = []
        for grp in self.properties:
            for code in grp.state_codes.split(", "):
                if code and code not in codes:
                    codes.append(code)
        return ", ".join(codes)

    @property
    def itinerary_days(self) -> list[tuple[str, list[PlotRow]]]:
        """Group plots by visit_date when available; otherwise by property order."""
        dated: dict[str, list[PlotRow]] = {}
        undated: list[PlotRow] = []
        for plot in self.all_plots:
            if plot.visit_date:
                dated.setdefault(plot.visit_date, []).append(plot)
            else:
                undated.append(plot)

        days: list[tuple[str, list[PlotRow]]] = []
        for date in sorted(dated):
            days.append((date, dated[date]))

        if undated:
            # Group undated by property as skeleton
            prop_map: dict[str, list[PlotRow]] = {}
            for plot in undated:
                prop_map.setdefault(plot.property, []).append(plot)
            for prop_name, plots in prop_map.items():
                days.append((f"[DATE TBD — {prop_name}]", plots))

        return days


# ---------------------------------------------------------------------------
# State-code to state-name mapping
# ---------------------------------------------------------------------------

_STATE_PREFIXES: dict[str, str] = {
    "NS": "NSW",
    "QD": "QLD",
    "SA": "SA",
    "WA": "WA",
    "NT": "NT",
    "VT": "VIC",
    "TS": "TAS",
    "AC": "ACT",
}


def _state_from_plot(plot_id: str) -> str:
    prefix = plot_id[:2].upper()
    return _STATE_PREFIXES.get(prefix, prefix)


# ---------------------------------------------------------------------------
# Database query
# ---------------------------------------------------------------------------


def _haversine_km(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    r = 6371.0
    dlat = math.radians(lat2 - lat1)
    dlon = math.radians(lon2 - lon1)
    a = (
        math.sin(dlat / 2) ** 2
        + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlon / 2) ** 2
    )
    return r * 2 * math.asin(math.sqrt(min(1.0, a)))


def query_trip_data(trip_id: str) -> TripData:
    """Pull all checklist fields for *trip_id* from campaigns.db + tern_plots.db."""
    if not CAMPAIGNS.exists():
        sys.exit(f"ERROR: campaigns.db not found at {CAMPAIGNS}. Run seed_campaigns.py first.")
    if not TERN_PLOTS.exists():
        sys.exit(f"ERROR: tern_plots.db not found at {TERN_PLOTS}.")

    con = sqlite3.connect(CAMPAIGNS)
    con.row_factory = sqlite3.Row
    con.execute(f"ATTACH DATABASE '{TERN_PLOTS.as_posix()}' AS tp")
    if ARD_STATE.exists():
        con.execute(f"ATTACH DATABASE '{ARD_STATE.as_posix()}' AS ard")

    # Campaign metadata
    row = con.execute(
        """
        SELECT trip_id, route_origin, route_dest, start_date, end_date,
               season, enso_phase, notes
        FROM campaigns
        WHERE trip_id = ?
        """,
        (trip_id,),
    ).fetchone()

    if row is None:
        # Try to fall back to the draft brief + kanban board (no DB row yet)
        con.close()
        return _trip_data_from_draft(trip_id)

    trip = TripData(
        trip_id=row["trip_id"],
        route_origin=row["route_origin"] or "",
        route_dest=row["route_dest"] or "",
        start_date=row["start_date"] or "",
        end_date=row["end_date"] or "",
        season=row["season"],
        enso_phase=row["enso_phase"],
        notes=row["notes"],
    )

    # Per-plot rows with TERN enrichment
    plots_rows = con.execute(
        """
        SELECT
            cp.plot,
            p.latitude,
            p.longitude,
            p.property,
            COALESCE(p.site_location_name, '') AS site_location_name,
            cp.visit_date,
            cp.access_notes,
            cp.collected,
            COALESCE(ib.ibra_bioregion_name, '') AS bioregion,
            COALESCE(nat.nvis_mvg_name, '')       AS mvg
        FROM campaign_plots cp
        JOIN tp.plots p ON p.plot = cp.plot
        LEFT JOIN tp.plot_ibra ib ON ib.plot = cp.plot
        LEFT JOIN tp.plot_nvis_national nat ON nat.plot = cp.plot
        WHERE cp.trip_id = ?
        ORDER BY p.property, cp.plot
        """,
        (trip_id,),
    ).fetchall()

    # Property metadata
    prop_rows = con.execute(
        """
        SELECT property_name, tenure, jurisdiction, access_status, email_address, notes
        FROM properties
        """,
    ).fetchall()
    prop_meta: dict[str, sqlite3.Row] = {r["property_name"]: r for r in prop_rows}

    # Group plots into PropertyGroup objects
    groups: dict[str, PropertyGroup] = {}
    for pr in plots_rows:
        pname = pr["property"] or "(Unknown)"
        if pname not in groups:
            meta = prop_meta.get(pname)
            groups[pname] = PropertyGroup(
                name=pname,
                tenure=meta["tenure"] if meta else None,
                jurisdiction=meta["jurisdiction"] if meta else None,
                access_status=meta["access_status"] if meta else None,
                email_address=meta["email_address"] if meta else None,
                notes=meta["notes"] if meta else None,
            )
        groups[pname].plots.append(
            PlotRow(
                plot=pr["plot"],
                latitude=pr["latitude"] or 0.0,
                longitude=pr["longitude"] or 0.0,
                property=pname,
                site_location_name=pr["site_location_name"],
                visit_date=pr["visit_date"],
                access_notes=pr["access_notes"],
                collected=pr["collected"],
                bioregion=pr["bioregion"] or None,
                mvg=pr["mvg"] or None,
                state=_state_from_plot(pr["plot"]),
            )
        )

    # Sort properties by uncollected count desc (mirrors shortlist ranking)
    trip.properties = sorted(groups.values(), key=lambda g: len(g.plots), reverse=True)

    try:
        meta_rows = con.execute(
            """
            SELECT visit_date, day_number, day_name, travel, property_visited, notes,
                   accommodation, accom_link, booking_status
            FROM itinerary_days
            WHERE trip_id = ?
            ORDER BY visit_date
            """,
            (trip_id,),
        ).fetchall()
        trip.itinerary_days_meta = [
            ItineraryDayMeta(
                visit_date=r["visit_date"],
                day_number=r["day_number"],
                day_name=r["day_name"] or "",
                travel=r["travel"] or "",
                property_visited=r["property_visited"] or "",
                notes=r["notes"] or "",
                accommodation=r["accommodation"] or "",
                accom_link=r["accom_link"] or "",
                booking_status=r["booking_status"] or "",
            )
            for r in meta_rows
        ]
    except sqlite3.OperationalError:
        trip.itinerary_days_meta = []

    con.close()
    return trip


def _trip_data_from_draft(trip_id: str) -> TripData:
    """
    Fall back to the draft brief markdown when no campaigns.db row exists yet.
    Parses docs/drafts/<trip-id>-draft.md and the matching kanban board.
    """
    from dronescape_planning.paths import BOARDS_DIR, DOCS_DRAFTS

    draft_path = DOCS_DRAFTS / f"{trip_id}-draft.md"
    if not draft_path.exists():
        sys.exit(
            f"ERROR: No DB row for '{trip_id}' in campaigns.db, and no draft brief at {draft_path}.\n"
            "Run: python scripts/shortlist.py --trip-id ... --draft to generate a draft first."
        )

    text = draft_path.read_text(encoding="utf-8")

    # Parse header
    route = ""
    field_window = ""
    m = re.search(r"^# Trip draft: \S+ \((.+)\)", text, re.MULTILINE)
    if m:
        route = m.group(1)
    m = re.search(r"^- Field window: (.+)", text, re.MULTILINE)
    if m:
        field_window = m.group(1).strip()

    # Dates
    start_date, end_date = "", ""
    m = re.search(r"(\d{4}-\d{2}-\d{2})\.\.(\d{4}-\d{2}-\d{2})", field_window)
    if m:
        start_date, end_date = m.group(1), m.group(2)

    # Season
    season = None
    m = re.search(r"\((winter|summer|autumn|spring|[^)]+)\)", field_window)
    if m:
        season = m.group(1)

    # Route origin / dest from trip_id pattern NN-ORIG-DEST-YYYY-MM
    parts = trip_id.split("-")
    route_origin, route_dest = "", ""
    if len(parts) >= 3:
        route_origin = parts[1]
        route_dest = parts[2]

    trip = TripData(
        trip_id=trip_id,
        route_origin=route_origin,
        route_dest=route_dest,
        start_date=start_date,
        end_date=end_date,
        season=season,
        enso_phase=None,
        notes=None,
    )

    # Parse property sections from the draft
    prop_blocks = re.split(r"^### \d+\. ", text, flags=re.MULTILINE)[1:]
    for block in prop_blocks:
        lines = block.strip().splitlines()
        header = lines[0]
        m = re.match(r"^(.+?) -- (\d+) plot", header)
        if not m:
            continue
        pname = m.group(1).strip()

        plot_ids: list[str] = []
        centroid_lat = centroid_lon = 0.0
        bioregion_str = mvg_str = tenure = jurisdiction = ""
        for line in lines[1:]:
            if line.startswith("- Plot IDs:"):
                ids_part = line.split(":", 1)[1].split("(")[0].strip()
                plot_ids = [x.strip() for x in ids_part.split(",") if x.strip()]
            elif line.startswith("- Centroid:"):
                coords = line.split(":", 1)[1].strip().split(",")
                if len(coords) == 2:
                    try:
                        centroid_lat = float(coords[0].strip())
                        centroid_lon = float(coords[1].strip())
                    except ValueError:
                        pass
            elif line.startswith("- Bioregion(s):"):
                bioregion_str = line.split(":", 1)[1].strip()
            elif line.startswith("- MVG(s):"):
                mvg_str = line.split(":", 1)[1].strip()
            elif line.startswith("- Tenure / jurisdiction:"):
                parts_tj = line.split(":", 1)[1].strip().split("/")
                tenure = parts_tj[0].strip()
                jurisdiction = parts_tj[1].strip() if len(parts_tj) > 1 else ""

        if not plot_ids:
            continue

        # Distribute centroid evenly across plot stubs (no real per-plot coords in draft)
        n = len(plot_ids)
        grp = PropertyGroup(
            name=pname,
            tenure=tenure or None,
            jurisdiction=jurisdiction or None,
            access_status=None,
            email_address=None,
            notes=None,
        )
        for i, pid in enumerate(plot_ids):
            # Jitter slightly so markers don't stack on the map
            jitter = 0.001 * (i - n / 2)
            grp.plots.append(
                PlotRow(
                    plot=pid,
                    latitude=centroid_lat + jitter,
                    longitude=centroid_lon + jitter,
                    property=pname,
                    site_location_name="",
                    visit_date=None,
                    access_notes=None,
                    collected=0,
                    bioregion=bioregion_str.split(",")[0].strip() if bioregion_str else None,
                    mvg=mvg_str.split(",")[0].strip() if mvg_str else None,
                    state=_state_from_plot(pid),
                )
            )
        trip.properties.append(grp)

    if not trip.properties:
        sys.exit(f"ERROR: Could not parse any properties from {draft_path}")

    return trip


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------

_TODO = "[TODO]"
_BOILERPLATE_COMMS = (
    "Site has reception: [ ] No  [ ] Yes: [ ] Full  [ ] Limited  (Check Mobile Coverage)\n\n"
    "Emergency communication equipment:\n"
    "- Garmin InReach Mini3\n"
    "- Personal Locator Beacon (PLB)\n"
    "- Satellite phone — TERN satellite phone 0147 160 278\n"
    "- VHF Radios / UHF Radios\n"
    "- Other: Whatsapp field-comms group daily once returned to camp; "
    "Phone call to Arko every 2 days between the hours 6–9 pm."
)
_BOILERPLATE_FIRST_AID = (
    "Standard first aid kits\n"
    "RFDS Medical Chest\n"
    "Personnel with remote first aid/CPR certification: Juan Carlos Montes Herrera; "
    f"{_TODO} (add all certified personnel)"
)
_BOILERPLATE_MEDICAL = (
    f"**Nearest RFDS base:** {_TODO}\n"
    f"**Clinic:** {_TODO}\n"
    "Phone: 000 (emergency retrieval) / (non-emergency — check Healthmap)\n"
    f"**Address:** {_TODO}\n"
    "Opening times: 24 hours (emergency)\n\n"
    "Check Healthmap to locate nearest services: https://www.healthmap.com.au"
)
_BOILERPLATE_WEATHER = (
    f"Location(s): see itinerary\n"
    "Check closest BOM station daily.\n"
    f"TFS reviewed: {_TODO}  Hotspots present: {_TODO}\n"
    f"Notes: {_TODO}"
)


def _md_row(label: str, value: str) -> str:
    value_clean = value.replace("\n", "<br>")
    return f"| **{label}** | {value_clean} |"


def _status_block(trip: TripData) -> str:
    confirmed = sum(1 for g in trip.properties if g.access_status == "confirmed")
    total = len(trip.properties)
    all_confirmed = confirmed == total and total > 0

    def tick(cond: bool) -> str:
        return "✅" if cond else "⬜"

    lines = [
        "**Overview**",
        f"- Planning: {tick(True)} commenced  {tick(bool(trip.start_date))} finalised",
        f"- Fieldwork locations: {tick(True)} identified  {tick(bool(trip.properties))} finalised",
        f"- Itinerary: {tick(bool(trip.start_date))} drafted  {tick(trip.itinerary_imported)} finalised",
        f"- Dates: {tick(bool(trip.start_date))} tentative  {tick(trip.itinerary_imported)} confirmed",
        "",
        "**Staffing and Logistics**",
        f"- Personnel: {tick(False)} tentative  {tick(False)} confirmed  {tick(False)} finalised",
        f"- Vehicles: {tick(False)} identified  {tick(False)} booked",
        "",
        "**Site Access**",
        f"- Scientific permits: {tick(False)} in progress  {tick(all_confirmed)} finalised",
        f"- Land access permissions: {tick(confirmed > 0)} commenced  {tick(all_confirmed)} confirmed",
        f"- Other access/permits: {_TODO}",
        "",
        "**Health and Safety**",
        f"- Risk assessment: {tick(False)} covered by current RA  {tick(False)} new RA applied  {tick(False)} approved",
        f"- AVCRM mission: {tick(False)} pending  {tick(False)} approved",
        f"- Site inductions: {tick(False)} commenced  {tick(False)} completed",
    ]
    return "\n".join(lines)


def _itinerary_block(trip: TripData) -> str:
    if trip.itinerary_days_meta:
        plots_by_date: dict[str, list[PlotRow]] = {}
        for date, plots in trip.itinerary_days:
            plots_by_date[date] = plots

        rows = [
            "| Date | Travel / Location | Plots | Notes |",
            "|------|-------------------|-------|-------|",
        ]
        for meta in trip.itinerary_days_meta:
            plots = plots_by_date.get(meta.visit_date, [])
            plot_ids = ", ".join(p.plot for p in plots) if plots else "—"
            location = meta.travel or meta.property_visited or meta.accommodation or "—"
            note_parts: list[str] = []
            if meta.notes:
                note_parts.append(meta.notes)
            if meta.accommodation:
                accom = meta.accommodation
                if meta.booking_status:
                    accom += f" ({meta.booking_status})"
                note_parts.append(f"Accom: {accom}")
            notes = "; ".join(note_parts) if note_parts else _TODO
            rows.append(f"| {meta.visit_date} | {location} | {plot_ids} | {notes} |")
        return "\n".join(rows)

    days = trip.itinerary_days
    if not days:
        return _TODO

    rows = ["| Date | Location / Property | Plots | Notes |", "|------|---------------------|-------|-------|"]
    for date, plots in days:
        props = sorted({p.property for p in plots})
        plot_ids = ", ".join(p.plot for p in plots)
        rows.append(f"| {date} | {'; '.join(props)} | {plot_ids} | {_TODO} |")
    return "\n".join(rows)


def _accommodation_block(trip: TripData) -> str:
    if not trip.itinerary_days_meta:
        return (
            f"Total nights away: {_TODO}\nNights bush camping: {_TODO}\n\n"
            "Location: See itinerary for daily accommodation."
        )
    lines: list[str] = []
    for meta in trip.itinerary_days_meta:
        if meta.accommodation:
            status = f" ({meta.booking_status})" if meta.booking_status else ""
            lines.append(f"{meta.visit_date}: {meta.accommodation}{status}")
    nights = len(lines)
    body = "\n".join(lines) if lines else "See daily itinerary table."
    return f"Total nights away: {nights}\nNights bush camping: {_TODO}\n\n{body}"


def _flight_areas_block(trip: TripData) -> str:
    rows = [
        "| Property | Plots (n) | Plot IDs | Centroid Lat | Centroid Lon | Bioregion | MVG | Tenure | Access |",
        "|----------|-----------|----------|-------------|-------------|-----------|-----|--------|--------|",
    ]
    for grp in trip.properties:
        rows.append(
            f"| {grp.name} | {len(grp.plots)} | {grp.plot_ids} "
            f"| {grp.centroid_lat:.4f} | {grp.centroid_lon:.4f} "
            f"| {grp.bioregions} | {grp.mvgs} "
            f"| {grp.tenure or _TODO} | {grp.access_status or _TODO} |"
        )
    return "\n".join(rows)


def _permits_block(trip: TripData) -> str:
    lines = [
        f"Valid state scientific permit: {_TODO} (permit authorisation date)",
        f"UAV permits required: {_TODO}",
        f"  - application submitted: {_TODO}  granted: {_TODO}",
        "",
        "**Landholders (from DB)**",
        "",
    ]
    for grp in trip.properties:
        status = grp.access_status or "unknown"
        lines.append(
            f"- **{grp.name}** · {grp.tenure or _TODO} · {grp.jurisdiction or _TODO} · "
            f"Access: {status} · Contact: {grp.email_address or _TODO}"
        )
    lines.append("")
    lines.append("Team to carry hard copies of permits and/or authorisations, as applicable.")
    return "\n".join(lines)


def _landholder_block(trip: TripData) -> str:
    lines: list[str] = []
    for grp in trip.properties:
        lines += [
            f"**{grp.name}**",
            f"- Tenure: {grp.tenure or _TODO}",
            f"- Jurisdiction: {grp.jurisdiction or _TODO}",
            f"- Access status: {grp.access_status or 'unknown'}",
            f"- Contact: {grp.email_address or _TODO}",
            f"- Notes: {grp.notes or _TODO}",
            f"- Access request: initiated {_TODO}  Granted: {_TODO}",
            "",
        ]
    return "\n".join(lines).strip()


def render_markdown(trip: TripData, maps_dir: Optional[Path] = None) -> str:
    """Return the full checklist as an Obsidian-compatible markdown string."""

    # Map embed helpers
    def embed_overview() -> str:
        if maps_dir:
            fname = f"{trip.trip_id}-overview.png"
            if (maps_dir / fname).exists():
                return f"\n\n![[{fname}]]\n"
        return f"\n\n_{_TODO}: paste overview map screenshot here_\n"

    def embed_property(grp: PropertyGroup) -> str:
        if maps_dir:
            safe = re.sub(r"[^\w\-]", "_", grp.name)[:40]
            fname = f"{trip.trip_id}-{safe}.png"
            if (maps_dir / fname).exists():
                return f"![[{fname}]]"
        return f"_{_TODO}: paste {grp.name} map screenshot_"

    season_str = f" · {trip.season}" if trip.season else ""
    enso_str = f" · ENSO: {trip.enso_phase}" if trip.enso_phase else ""
    date_range = (
        f"{trip.start_date} → {trip.end_date}"
        if trip.start_date and trip.end_date
        else _TODO
    )

    lines = [
        f"# Trip Logistics Checklist",
        f"## DroneScape: Fieldwork — {trip.trip_id}",
        f"",
        f"> **Route:** {trip.route_origin} → {trip.route_dest}{season_str}{enso_str}",
        f"> **Dates:** {date_range}",
        f"> **Properties:** {len(trip.properties)} · **Plots:** {len(trip.all_plots)}",
        f"",
        f"---",
        f"",
        f"## Status Summary",
        f"",
        _status_block(trip),
        f"",
        f"---",
        f"",
        f"## Outstanding Issues",
        f"",
        f"> {_TODO}: Provide a detailed description of any issues yet to be resolved.",
        f"",
        f"---",
        f"",
        f"## Overview",
        f"",
        f"| Field | Value |",
        f"|-------|-------|",
        _md_row("Personnel", f"Person 1: {_TODO} · Person 2: {_TODO}\nEmergency Contact: {_TODO}"),
        _md_row("Location(s)", trip.location_names or _TODO),
        _md_row("State(s)", trip.states or _TODO),
        _md_row("Closest town(s)", _TODO),
        _md_row("Departure date", trip.start_date or _TODO),
        _md_row("Return date", trip.end_date or _TODO),
        _md_row("Team departing from", trip.route_origin or _TODO),
        _md_row("Work beginning location", _TODO),
        _md_row("Work ending location", _TODO),
        _md_row("Team returning to", trip.route_dest or _TODO),
        f"",
        f"### Site Activity / Itinerary",
        f"",
        _itinerary_block(trip),
        f"",
        f"### Site Map / Flight Areas",
        embed_overview(),
        f"",
        _flight_areas_block(trip),
        f"",
    ]

    # Per-property map embeds
    for grp in trip.properties:
        lines += [
            f"#### {grp.name} ({len(grp.plots)} plots)",
            embed_property(grp),
            f"",
        ]

    lines += [
        f"---",
        f"",
        f"## Staffing and Logistics",
        f"",
        f"| Field | Value |",
        f"|-------|-------|",
        _md_row("Personnel", f"Person 1: {_TODO} Phone: {_TODO}\nPerson 2: {_TODO} Phone: {_TODO}\nVolunteer: {_TODO}"),
        _md_row("Communications", _BOILERPLATE_COMMS),
        _md_row(
            "Vehicles and trailers",
            f"Vehicles: {_TODO} (Yes/No — vehicle name, make/model)\nNotes: {_TODO}\n\nTrailers: {_TODO}\nNotes: {_TODO}",
        ),
        _md_row(
            "Intended route",
            "\n".join(
                f"Day {i + 1}: {meta.visit_date} — {meta.travel or meta.property_visited or '—'}"
                for i, meta in enumerate(trip.itinerary_days_meta)
            )
            if trip.itinerary_days_meta
            else (
                "\n".join(
                    f"Day {i + 1}: {date} — {', '.join(sorted({p.property for p in plots}))}"
                    for i, (date, plots) in enumerate(trip.itinerary_days)
                )
                if trip.itinerary_days
                else _TODO
            ),
        ),
        _md_row("Accommodation", _accommodation_block(trip)),
        _md_row("First Aid", _BOILERPLATE_FIRST_AID),
        _md_row("Other special equipment", f"Notes: {_TODO}"),
        f"",
        f"---",
        f"",
        f"## Site Access and Biosecurity",
        f"",
        f"| Field | Value |",
        f"|-------|-------|",
        _md_row("Permits", _permits_block(trip)),
        _md_row("Landholder access permissions", _landholder_block(trip)),
        f"",
        f"---",
        f"",
        f"## Health and Safety",
        f"",
        f"| Field | Value |",
        f"|-------|-------|",
        _md_row("Medical and emergency services", _BOILERPLATE_MEDICAL),
        _md_row("Weather and fire", _BOILERPLATE_WEATHER),
        _md_row("Site-specific trainings", f"Site-specific inductions required: {_TODO}\nNotes: {_TODO}"),
        f"",
        f"---",
        f"",
        f"## Site Maps",
        f"",
        f"_Paste or embed screenshot images of GPS tracks, flight areas, and route maps here._",
        f"",
        embed_overview(),
        f"",
    ]

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

def _cell_text(table, row_idx: int, col_idx: int) -> str:
    return table.rows[row_idx].cells[col_idx].text


def _set_run_black(run) -> None:
    """Force run text to black (override Word theme/hyperlink purple)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    for tag in ("w:themeColor", "w:themeTint", "w:themeShade", "w:themeFill"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), "000000")


def _set_cell(
    table,
    row_idx: int,
    col_idx: int,
    text: str,
    images: Optional[list[Path]] = None,
    image_width_inches: float = 5.5,
) -> None:
    """Replace cell content with black text and optional embedded images."""
    from docx.oxml.ns import qn
    from docx.shared import Inches

    cell = table.rows[row_idx].cells[col_idx]
    tc = cell._tc

    for child in list(tc):
        if child.tag in (qn("w:p"), qn("w:tbl"), qn("w:sdt")):
            tc.remove(child)

    for line in text.split("\n"):
        para = cell.add_paragraph()
        run = para.add_run(line)
        _set_run_black(run)

    if images:
        for img_path in images:
            if not img_path.exists():
                continue
            para = cell.add_paragraph()
            run = para.add_run()
            run.add_picture(str(img_path), width=Inches(image_width_inches))
            caption = cell.add_paragraph()
            cap_run = caption.add_run(img_path.stem.replace("-", " "))
            _set_run_black(cap_run)


def _trip_map_paths(trip: TripData, maps_dir: Path) -> tuple[Optional[Path], list[Path]]:
    """Return (overview_path, per_property_paths) if PNGs exist."""
    overview = maps_dir / f"{trip.trip_id}-overview.png"
    overview_path = overview if overview.exists() else None

    property_paths: list[Path] = []
    for grp in trip.properties:
        safe = re.sub(r"[^\w\-]", "_", grp.name)[:40]
        path = maps_dir / f"{trip.trip_id}-{safe}.png"
        if path.exists():
            property_paths.append(path)

    return overview_path, property_paths


def _embed_site_maps_section(doc, trip: TripData, maps_dir: Path) -> None:
    """Add overview + per-property map images under the 'Site maps' heading."""
    from docx.oxml import OxmlElement
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph

    overview_path, _ = _trip_map_paths(trip, maps_dir)

    all_images: list[tuple[str, Path]] = []
    if overview_path:
        all_images.append(("Overview — all sites", overview_path))
    for grp in trip.properties:
        safe = re.sub(r"[^\w\-]", "_", grp.name)[:40]
        path = maps_dir / f"{trip.trip_id}-{safe}.png"
        if path.exists():
            all_images.append((grp.name, path))

    if not all_images:
        return

    site_maps_para = None
    for para in doc.paragraphs:
        if para.text.strip() == "Site maps":
            site_maps_para = para
            break

    if site_maps_para is None:
        return

    def insert_after(ref_para, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        ref_para._element.addnext(new_p)
        new_para = Paragraph(new_p, ref_para._parent)
        if text:
            run = new_para.add_run(text)
            _set_run_black(run)
        return new_para

    anchor = site_maps_para
    for title, img_path in all_images:
        title_para = insert_after(anchor, title)
        title_para.runs[0].bold = True
        anchor = title_para

        img_para = insert_after(anchor)
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(6.0))
        anchor = img_para


def render_docx(
    trip: TripData,
    template_path: Path,
    output_path: Path,
    maps_dir: Optional[Path] = None,
) -> None:
    """Clone the DOCX template and fill in trip data."""
    try:
        from docx import Document
    except ImportError:
        sys.exit("ERROR: python-docx is not installed. Run: pip install -e '.[docx]'")

    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    table = doc.tables[0]
    overview_img: Optional[Path] = None
    if maps_dir is not None:
        overview_img, _ = _trip_map_paths(trip, maps_dir)

    for para in doc.paragraphs:
        if "Fieldwork" in para.text and "Dates From" in para.text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"DroneScape: Fieldwork — {trip.trip_id}"
                _set_run_black(para.runs[0])
            else:
                run = para.add_run(f"DroneScape: Fieldwork — {trip.trip_id}")
                _set_run_black(run)
            break

    season_str = f" · {trip.season}" if trip.season else ""
    enso_str = f" · ENSO: {trip.enso_phase}" if trip.enso_phase else ""
    date_range = (
        f"{trip.start_date} to {trip.end_date}"
        if trip.start_date and trip.end_date
        else _TODO
    )

    # Row 1 (Status summary value)
    _set_cell(table, 1, 1, _status_block(trip).replace("✅", "[x]").replace("⬜", "[ ]"))

    # Row 3 (Outstanding issues)
    _set_cell(table, 3, 1, f"{_TODO}: Provide a detailed description of any issues yet to be resolved.")

    # Row 4 (Overview header — leave label, patch value)
    _set_cell(
        table, 4, 1,
        f"OVERVIEW\nPerson 1: {_TODO}\tPerson 2: {_TODO}\nEmergency Contact: {_TODO}"
    )

    # Row 5: Location(s)
    _set_cell(table, 5, 1, trip.location_names or _TODO)

    # Row 6: State(s)
    _set_cell(table, 6, 1, trip.states or _TODO)

    # Row 7: Closest town(s)
    _set_cell(table, 7, 1, _TODO)

    # Row 8: Departure date + personnel note
    _set_cell(
        table, 8, 1,
        f"{trip.start_date or _TODO}\n"
        f"All staff attending entire field trip: {_TODO}\n"
        f"All volunteers: {_TODO}\n"
        f"Note any mid-trip personnel changes."
    )

    # Row 9: Return date
    _set_cell(table, 9, 1, trip.end_date or _TODO)

    # Row 10: Team departing from
    _set_cell(table, 10, 1, trip.route_origin or _TODO)

    # Row 11: Work beginning location
    _set_cell(table, 11, 1, _TODO)

    # Row 12: Work ending location
    _set_cell(table, 12, 1, _TODO)

    # Row 13: Team returning to
    _set_cell(table, 13, 1, trip.route_dest or _TODO)

    # Row 14: Site activity
    _set_cell(table, 14, 1, _itinerary_as_text(trip))

    # Row 15: Site map / flight areas (+ overview map if available)
    row15_images = [overview_img] if overview_img else None
    _set_cell(
        table,
        15,
        1,
        _flight_areas_as_text(trip),
        images=row15_images,
        image_width_inches=5.0,
    )

    # Row 17: Personnel
    _set_cell(table, 17, 1, f"Person 1: {_TODO}\tPhone: {_TODO}\nPerson 2: {_TODO}\tPhone: {_TODO}\nVolunteer: {_TODO}")

    # Row 18: Communications (keep boilerplate)
    _set_cell(table, 18, 1, _BOILERPLATE_COMMS.replace("\n\n", "\n").replace("[ ] ", ""))

    # Row 19: Vehicles
    _set_cell(
        table, 19, 1,
        f"Vehicles: {_TODO} (Yes — vehicle name, make/model / No)\nNotes: {_TODO}\n\nTrailers: {_TODO}\nNotes: {_TODO}"
    )

    # Row 20: Intended route
    route_text = "Screenshots of tracks and roads for each day of travel\n"
    for i, (date, plots) in enumerate(trip.itinerary_days, 1):
        props = sorted({p.property for p in plots})
        route_text += f"Day {i}: {date} — {', '.join(props)}\n"
    _set_cell(table, 20, 1, route_text.strip() or _TODO)

    # Row 21: Accommodation
    _set_cell(
        table, 21, 1,
        f"Total number of nights away: {_TODO}\n"
        f"Number of nights bush camping: {_TODO}\n\n"
        "Location: See itinerary for daily accommodation name and locations."
    )

    # Row 22: First Aid (keep boilerplate)
    _set_cell(table, 22, 1, _BOILERPLATE_FIRST_AID)

    # Row 23: Other special equipment
    _set_cell(table, 23, 1, f"Notes: {_TODO}")

    # Row 25: Permits
    _set_cell(table, 25, 1, _permits_as_text(trip))

    # Row 26: Landholder access
    _set_cell(table, 26, 1, _landholder_as_text(trip))

    # Row 28: Medical and emergency services
    _set_cell(table, 28, 1, _BOILERPLATE_MEDICAL.replace("**", ""))

    # Row 29: Weather and fire
    _set_cell(table, 29, 1, _BOILERPLATE_WEATHER)

    # Row 30: Site inductions
    _set_cell(table, 30, 1, f"Site-specific inductions required: {_TODO}\nNotes: {_TODO}")

    if maps_dir is not None:
        _embed_site_maps_section(doc, trip, maps_dir)

    doc.save(output_path)


def _itinerary_as_text(trip: TripData) -> str:
    if trip.itinerary_days_meta:
        lines = []
        plots_by_date = dict(trip.itinerary_days)
        for meta in trip.itinerary_days_meta:
            plots = plots_by_date.get(meta.visit_date, [])
            plot_ids = ", ".join(p.plot for p in plots) if plots else "—"
            location = meta.travel or meta.property_visited or "—"
            extra = f" — {meta.notes}" if meta.notes else ""
            lines.append(f"{meta.visit_date}: {location} — Plots: {plot_ids}{extra}")
        return "\n".join(lines)

    days = trip.itinerary_days
    if not days:
        return _TODO
    lines = []
    for date, plots in days:
        props = sorted({p.property for p in plots})
        plot_ids = ", ".join(p.plot for p in plots)
        lines.append(f"{date}: {'; '.join(props)} — Plots: {plot_ids}")
    return "\n".join(lines)


def _flight_areas_as_text(trip: TripData) -> str:
    lines = ["Property | Plots | Centroid Lat, Lon | Tenure | Access"]
    for grp in trip.properties:
        lines.append(
            f"{grp.name} | {len(grp.plots)} | {grp.centroid_lat:.4f}, {grp.centroid_lon:.4f} "
            f"| {grp.tenure or _TODO} | {grp.access_status or _TODO}"
        )
    return "\n".join(lines)


def _permits_as_text(trip: TripData) -> str:
    lines = [
        f"Valid state scientific permit: {_TODO}",
        f"UAV permits required: {_TODO}  application submitted: {_TODO}  granted: {_TODO}",
        "",
        "Landholders:",
    ]
    for grp in trip.properties:
        lines.append(
            f"  {grp.name}: {grp.tenure or _TODO} / {grp.jurisdiction or _TODO} — "
            f"access: {grp.access_status or 'unknown'}"
        )
    lines.append("")
    lines.append("Team to carry hard copies of permits and/or authorisations, as applicable.")
    return "\n".join(lines)


def _landholder_as_text(trip: TripData) -> str:
    lines: list[str] = []
    for grp in trip.properties:
        lines += [
            f"{grp.name}",
            f"  Tenure: {grp.tenure or _TODO}  Jurisdiction: {grp.jurisdiction or _TODO}",
            f"  Access status: {grp.access_status or 'unknown'}",
            f"  Contact: {grp.email_address or _TODO}",
            f"  Access request: Initiated {_TODO}  Granted: {_TODO}",
            "",
        ]
    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Map generator
# ---------------------------------------------------------------------------

def generate_maps(trip: TripData, maps_dir: Path) -> list[Path]:
    """
    Produce static PNG maps using geopandas + contextily.
    Falls back to a folium HTML map if contextily tile fetch fails.

    Returns a list of paths to generated files.
    """
    try:
        import contextily as ctx
        import geopandas as gpd
        import matplotlib.pyplot as plt
        from shapely.geometry import Point
    except ImportError:
        sys.exit(
            "ERROR: Mapping libraries not installed. Run: pip install -e '.[maps]'"
        )

    maps_dir.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []

    plots = trip.all_plots
    if not plots:
        print("WARNING: No plots with coordinates — skipping map generation.")
        return generated

    def make_gdf(plot_list: list[PlotRow]) -> "gpd.GeoDataFrame":
        import pandas as pd
        data = {
            "plot": [p.plot for p in plot_list],
            "property": [p.property for p in plot_list],
            "geometry": [Point(p.longitude, p.latitude) for p in plot_list],
        }
        gdf = gpd.GeoDataFrame(data, crs="EPSG:4326")
        return gdf.to_crs(epsg=3857)  # Web Mercator for contextily

    def save_map(gdf: "gpd.GeoDataFrame", title: str, out_path: Path, zoom: Optional[int] = None) -> Path:
        fig, ax = plt.subplots(figsize=(12, 9))
        gdf.plot(ax=ax, color="red", markersize=60, zorder=5, alpha=0.85)

        # Label each plot
        for _, row in gdf.iterrows():
            ax.annotate(
                row["plot"],
                xy=(row.geometry.x, row.geometry.y),
                xytext=(6, 6),
                textcoords="offset points",
                fontsize=7,
                color="darkred",
                zorder=6,
            )

        # Pad bounds by 15%
        bounds = gdf.total_bounds  # minx, miny, maxx, maxy
        xpad = max((bounds[2] - bounds[0]) * 0.15, 5000)
        ypad = max((bounds[3] - bounds[1]) * 0.15, 5000)
        ax.set_xlim(bounds[0] - xpad, bounds[2] + xpad)
        ax.set_ylim(bounds[1] - ypad, bounds[3] + ypad)

        try:
            ctx.add_basemap(ax, source=ctx.providers.OpenStreetMap.Mapnik, zoom=zoom or "auto")
        except Exception as e:
            print(f"  WARNING: Basemap fetch failed ({e}). Saving without basemap.")

        ax.set_axis_off()
        ax.set_title(title, fontsize=11, pad=8)
        fig.tight_layout()
        fig.savefig(out_path, dpi=150, bbox_inches="tight")
        plt.close(fig)
        return out_path

    # Overview map — all plots
    print(f"  Generating overview map ({len(plots)} plots)...")
    gdf_all = make_gdf(plots)
    overview_path = maps_dir / f"{trip.trip_id}-overview.png"
    save_map(gdf_all, f"{trip.trip_id} — All Sites Overview", overview_path)
    generated.append(overview_path)
    print(f"    Saved: {overview_path.name}")

    # Per-property maps
    for grp in trip.properties:
        if len(grp.plots) == 0:
            continue
        safe = re.sub(r"[^\w\-]", "_", grp.name)[:40]
        out_path = maps_dir / f"{trip.trip_id}-{safe}.png"
        print(f"  Generating map for {grp.name} ({len(grp.plots)} plots)...")
        gdf_prop = make_gdf(grp.plots)
        save_map(gdf_prop, f"{grp.name} — {len(grp.plots)} plots", out_path)
        generated.append(out_path)
        print(f"    Saved: {out_path.name}")

    return generated


# ---------------------------------------------------------------------------
# Folium interactive map (bonus — for Obsidian HTML embed)
# ---------------------------------------------------------------------------

def generate_folium_map(trip: TripData, maps_dir: Path) -> Optional[Path]:
    """Generate an interactive folium HTML map alongside the static PNGs."""
    try:
        import folium
        from folium.plugins import MarkerCluster
    except ImportError:
        return None

    plots = trip.all_plots
    if not plots:
        return None

    centre_lat = sum(p.latitude for p in plots) / len(plots)
    centre_lon = sum(p.longitude for p in plots) / len(plots)

    m = folium.Map(location=[centre_lat, centre_lon], zoom_start=7, tiles="OpenStreetMap")
    cluster = MarkerCluster().add_to(m)

    colors = ["red", "blue", "green", "purple", "orange", "darkred", "darkblue", "darkgreen"]
    color_map: dict[str, str] = {}
    for i, grp in enumerate(trip.properties):
        color_map[grp.name] = colors[i % len(colors)]

    for p in plots:
        color = color_map.get(p.property, "gray")
        popup_html = (
            f"<b>{p.plot}</b><br>"
            f"Property: {p.property}<br>"
            f"Lat: {p.latitude:.5f}, Lon: {p.longitude:.5f}<br>"
            f"Bioregion: {p.bioregion or '—'}<br>"
            f"MVG: {p.mvg or '—'}<br>"
            f"Collected: {'Yes' if p.collected else 'No'}"
        )
        folium.Marker(
            location=[p.latitude, p.longitude],
            popup=folium.Popup(popup_html, max_width=280),
            tooltip=p.plot,
            icon=folium.Icon(color=color, icon="circle", prefix="fa"),
        ).add_to(cluster)

    out_path = maps_dir / f"{trip.trip_id}-interactive.html"
    m.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Generate a Trip Logistics Checklist from the DroneScape databases.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument("--trip-id", required=True, metavar="ID",
                   help="Trip identifier, e.g. 05-DRW-TOP-2027-06.")
    p.add_argument("--docx", action="store_true",
                   help="Also generate a filled DOCX from the template.")
    p.add_argument("--template", type=Path, default=None, metavar="PATH",
                   help=f"Path to the DOCX template (default: {TEMPLATE_DOCX}).")
    p.add_argument("--maps", action="store_true",
                   help="Generate static PNG site maps via geopandas + contextily.")
    p.add_argument("--folium", action="store_true",
                   help="Also generate an interactive folium HTML map.")
    p.add_argument("--out-dir", type=Path, default=None, metavar="DIR",
                   help=f"Output directory (default: {DOCS_CHECKLISTS}).")
    return p.parse_args()


def main() -> None:
    args = parse_args()
    trip_id = args.trip_id
    out_dir = args.out_dir or DOCS_CHECKLISTS
    maps_dir = out_dir / "maps"
    template_path = args.template or TEMPLATE_DOCX

    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"Querying trip data for {trip_id}...")
    trip = query_trip_data(trip_id)
    print(
        f"  Found: {len(trip.properties)} properties, "
        f"{len(trip.all_plots)} plots, "
        f"{trip.start_date or '?'} -> {trip.end_date or '?'}"
    )

    # Maps first (so they can be embedded in markdown)
    if args.maps or args.folium:
        maps_dir.mkdir(parents=True, exist_ok=True)

    if args.maps:
        print("Generating site maps...")
        generate_maps(trip, maps_dir)

    if args.folium:
        print("Generating interactive folium map...")
        html_path = generate_folium_map(trip, maps_dir)
        if html_path:
            print(f"  Saved: {html_path.name}")

    # Markdown
    md_path = out_dir / f"{trip_id}-checklist.md"
    print(f"Writing markdown checklist to {md_path}...")
    md_content = render_markdown(trip, maps_dir=maps_dir if args.maps else None)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"  Done: {md_path}")

    # DOCX
    if args.docx:
        if not template_path.exists():
            print(
                f"WARNING: DOCX template not found at {template_path}.\n"
                "  Copy the template to the default location or pass --template PATH.\n"
                "  Skipping DOCX output."
            )
        else:
            docx_path = out_dir / f"{trip_id}-checklist.docx"
            print(f"Writing DOCX checklist to {docx_path}...")
            docx_maps_dir = maps_dir if maps_dir.exists() and any(maps_dir.glob("*.png")) else None
            if args.docx and docx_maps_dir is None and not args.maps:
                print(
                    "  NOTE: No map PNGs found. Re-run with --maps to embed site maps in the DOCX."
                )
            render_docx(trip, template_path, docx_path, maps_dir=docx_maps_dir)
            print(f"  Done: {docx_path}")

    print("\nChecklist generation complete.")
    print(f"  Output directory: {out_dir}")


if __name__ == "__main__":
    main()
